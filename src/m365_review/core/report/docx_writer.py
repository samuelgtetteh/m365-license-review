"""Word report writer (python-docx).

A client-facing narrative document that leads with the license optimization
summary, then details each finding, then the full license inventory. Formatting
is clean and professional: styled headings, a shaded table header, and USD
formatting. This is the document an MSP hands to a client.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from m365_review.core.models import AuditResult, Severity

_NAVY = RGBColor(0x1F, 0x38, 0x64)
_HEADER_SHADE = "1F3864"
_GREY = RGBColor(0x59, 0x59, 0x59)


def _money(value: float | None) -> str:
    return "—" if value is None else f"${value:,.2f}"


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _style_header_row(table) -> None:
    for cell in table.rows[0].cells:
        _shade_cell(cell, _HEADER_SHADE)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _inventory_table(doc, title: str, rows, *, show_costs: bool) -> None:
    """Render one inventory category (paid or free) as a heading + table."""
    doc.add_heading(title, level=2)
    if not rows:
        doc.add_paragraph("(none)")
        return
    cols = ["Product", "Purchased", "Assigned", "Unused", "Unit $/mo", "Unused $/mo"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    for i, text in enumerate(cols):
        t.rows[0].cells[i].paragraphs[0].add_run(text)
    _style_header_row(t)
    for r in rows:
        c = t.add_row().cells
        c[0].text = r.display_name
        c[1].text = str(r.purchased)
        c[2].text = str(r.assigned)
        c[3].text = str(r.slack)
        if show_costs:
            c[4].text = _money(r.unit_price_usd)
            c[5].text = _money(r.slack_monthly_usd)
        else:
            c[4].text = "Free"
            c[5].text = "—"


def write_docx(result: AuditResult, path: Path) -> Path:
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Microsoft 365 License Optimization Report", level=0)
    title.runs[0].font.color.rgb = _NAVY
    sub = doc.add_paragraph()
    sub.add_run(
        f"{result.tenant_display_name}  ·  {result.report_date.isoformat()}"
    ).italic = True

    # --- Tenant facts ---
    facts = doc.add_paragraph()
    facts.add_run("Tenant ID: ").bold = True
    facts.add_run(result.tenant_id)
    r = facts.add_run("\nReport is read-only — no changes were made to the tenant.")
    r.font.color.rgb = _GREY
    r.font.size = Pt(9)

    # --- How to read ---
    doc.add_heading("How to read this report", level=1)
    doc.add_paragraph(
        "This report compares the Microsoft 365 licenses your tenant owns against how they are "
        "actually used, and recommends specific, prioritized actions to reduce cost. Savings "
        "figures are estimates based on the price list in effect when the report was generated; "
        "please confirm each recommendation before acting."
    )

    # --- Optimization summary (the headline) ---
    doc.add_heading("License optimization summary", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        f"We identified {len(result.findings)} optimization finding(s) with an estimated "
        f"{_money(result.total_monthly_savings_usd)} per month "
        f"({_money(result.total_annual_savings_usd)} per year) in potential savings."
    )
    run.bold = True

    sev = result.severity_counts()
    doc.add_paragraph(
        f"Severity breakdown — High: {sev['high']}, Medium: {sev['medium']}, "
        f"Low: {sev['low']}, Info: {sev['info']}.",
    )

    # Prioritized recommendations table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(["#", "Severity", "Recommendation", "Users", "Monthly $"]):
        hdr[i].paragraphs[0].add_run(text)
    _style_header_row(table)
    for i, f in enumerate(result.findings, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = f.severity.value.capitalize()
        cells[2].text = f"{f.title} — {f.recommendation}"
        cells[3].text = str(f.affected_count)
        cells[4].text = _money(f.estimated_monthly_savings_usd)
    totals = table.add_row().cells
    totals[2].paragraphs[0].add_run("TOTAL").bold = True
    totals[4].paragraphs[0].add_run(_money(result.total_monthly_savings_usd)).bold = True

    # --- Detailed findings ---
    doc.add_heading("Findings in detail", level=1)
    if not result.findings:
        doc.add_paragraph("No optimization findings — this tenant's licensing looks well-managed.")
    for f in result.findings:
        h = doc.add_heading(f"{f.rule_id} · {f.title}", level=2)
        h.runs[0].font.color.rgb = _NAVY
        meta = doc.add_paragraph()
        meta.add_run(f"Severity: {f.severity.value.capitalize()}   ").bold = True
        meta.add_run(
            f"Estimated savings: {_money(f.estimated_monthly_savings_usd)}/mo "
            f"({_money(f.estimated_annual_savings_usd)}/yr)"
        )
        doc.add_paragraph(f.description)
        rec = doc.add_paragraph()
        rec.add_run("Recommendation: ").bold = True
        rec.add_run(f.recommendation)

        if f.affected_users:
            doc.add_paragraph("Affected users:", style="Intense Quote")
            ut = doc.add_table(rows=1, cols=3)
            ut.style = "Light List Accent 1"
            for i, text in enumerate(["User", "Detail", "Monthly $"]):
                ut.rows[0].cells[i].paragraphs[0].add_run(text)
            _style_header_row(ut)
            for u in f.affected_users:
                c = ut.add_row().cells
                c[0].text = u.user_principal_name or (u.display_name or "—")
                c[1].text = u.detail or ""
                c[2].text = _money(u.monthly_cost_usd)
        for note in f.notes:
            n = doc.add_paragraph(f"Note: {note}")
            n.runs[0].font.color.rgb = _GREY
            n.runs[0].font.size = Pt(9)

    # --- License inventory (split into paid vs free categories) ---
    doc.add_heading("License inventory", level=1)
    paid = [r for r in result.inventory if not r.is_unlimited]
    free = [r for r in result.inventory if r.is_unlimited]

    _inventory_table(doc, "Paid licenses", paid, show_costs=True)
    if free:
        _inventory_table(
            doc, "Free / self-service licenses (excluded from totals)", free, show_costs=False
        )

    # --- Caveats ---
    if result.caveats:
        doc.add_heading("Notes & data caveats", level=1)
        for cav in result.caveats:
            doc.add_paragraph(cav, style="List Bullet")

    doc.save(path)
    return path
